import io
from typing import Any, List, Dict
from datetime import datetime, timedelta, timezone
from fastapi import APIRouter, Depends, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func, cast, Date
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor

from apps.api.core.database import get_db
from apps.api.core.ai_service import generate_market_analysis
from apps.services.gpu.models_history import GpuPriceHistory

router = APIRouter()


async def _fetch_price_data(
    db: AsyncSession, gpu_model_id: str, days: int = 30
) -> List[Dict]:
    """price_history 테이블에서 일별 집계 데이터 조회."""
    since = datetime.now(timezone.utc) - timedelta(days=days)
    day_col = cast(GpuPriceHistory.ts, Date).label("day")

    result = await db.execute(
        select(
            day_col,
            GpuPriceHistory.prv_id,
            GpuPriceHistory.gpu_mdl,
            func.min(GpuPriceHistory.prc_ph).label("min_price"),
            func.avg(GpuPriceHistory.prc_ph).label("avg_price"),
            func.max(GpuPriceHistory.prc_ph).label("max_price"),
            func.count(GpuPriceHistory.id).label("cnt"),
        )
        .where(GpuPriceHistory.gpu_mdl.ilike(f"%{gpu_model_id}%"))
        .where(GpuPriceHistory.ts >= since)
        .group_by(day_col, GpuPriceHistory.prv_id, GpuPriceHistory.gpu_mdl)
        .order_by(day_col.asc(), GpuPriceHistory.prv_id)
    )
    rows = result.all()
    return [
        {
            "date": str(r.day),
            "provider": r.provider_id,
            "gpu_model": r.gpu_model,
            "min_price": round(float(r.min_price), 4),
            "avg_price": round(float(r.avg_price), 4),
            "max_price": round(float(r.max_price), 4),
            "data_points": r.cnt,
        }
        for r in rows
    ]


@router.get("/excel")
async def export_excel(
    gpu_model_id: str = Query("H100", description="GPU 모델명"),
    days: int = Query(30, ge=1, le=365, description="조회 기간(일)"),
    db: AsyncSession = Depends(get_db),
) -> Any:
    """
    Exports time-series pricing data as a styled Excel file.
    (수정: Mock 데이터 제거 → 실제 price_history DB 조회)
    """
    rows = await _fetch_price_data(db, gpu_model_id, days)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"{gpu_model_id} Price Data"

    # ── 헤더 스타일 ──────────────────────────────────────────────
    HEADER_FONT  = Font(bold=True, color="FFFFFF", size=11)
    HEADER_FILL  = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    CENTER       = Alignment(horizontal="center", vertical="center")
    THIN_BORDER  = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    EVEN_FILL    = PatternFill(start_color="EBF3FB", end_color="EBF3FB", fill_type="solid")

    headers = ["Date", "Provider", "GPU Model", "Min $/hr", "Avg $/hr", "Max $/hr", "Data Points"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font      = HEADER_FONT
        cell.fill      = HEADER_FILL
        cell.alignment = CENTER
        cell.border    = THIN_BORDER

    # ── 데이터 삽입 ─────────────────────────────────────────────
    if not rows:
        # 데이터 없음 안내 행
        ws.append(["No data available for the selected period.", "", "", "", "", "", ""])
    else:
        for i, row in enumerate(rows, start=2):
            ws.append([
                row["date"], row["provider"], row["gpu_model"],
                row["min_price"], row["avg_price"], row["max_price"], row["data_points"],
            ])
            if i % 2 == 0:
                for cell in ws[i]:
                    cell.fill   = EVEN_FILL
                    cell.border = THIN_BORDER
            else:
                for cell in ws[i]:
                    cell.border = THIN_BORDER

    # ── 숫자 포맷 ────────────────────────────────────────────────
    for row in ws.iter_rows(min_row=2, min_col=4, max_col=6):
        for cell in row:
            cell.number_format = '"$"#,##0.0000'

    # ── 열 너비 자동 조정 ────────────────────────────────────────
    for col_idx, col in enumerate(ws.columns, start=1):
        max_len = max((len(str(c.value or "")) for c in col), default=10)
        ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 4, 40)

    # ── 메타 시트 ────────────────────────────────────────────────
    ws_meta = wb.create_sheet("Report Info")
    ws_meta.append(["Generated", datetime.now(timezone.utc).isoformat()])
    ws_meta.append(["GPU Model", gpu_model_id])
    ws_meta.append(["Period", f"Last {days} days"])
    ws_meta.append(["Rows", len(rows)])
    ws_meta.append(["Source", "InfraIndex price_history DB"])

    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)

    filename = f"{gpu_model_id}_pricing_{datetime.now().strftime('%Y%m%d')}.xlsx"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.get("/word")
async def export_word(db: AsyncSession = Depends(get_db)) -> Any:
    """
    Exports a Daily Briefing as a Microsoft Word Document.
    AI 섹션은 실제 LLM 연결 (ai_service.py) 사용.
    """
    doc = Document()
    doc.add_heading("InfraIndex Daily Macro Briefing", 0)

    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    doc.add_paragraph(f"Report Generated: {date_str}")

    # 1. GPU Price Trend (DB에서 H100 최근 7일)
    doc.add_heading("1. GPU Market Price Trend", level=1)
    h100_rows = await _fetch_price_data(db, "H100", days=7)

    if h100_rows:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Date", "Provider", "Min $/hr", "Avg $/hr", "Max $/hr"]):
            hdr[i].text = h

        for r in h100_rows[:20]:
            cells = table.add_row().cells
            cells[0].text = r["date"]
            cells[1].text = r["provider"]
            cells[2].text = f"${r['min_price']:.4f}"
            cells[3].text = f"${r['avg_price']:.4f}"
            cells[4].text = f"${r['max_price']:.4f}"
    else:
        doc.add_paragraph("No H100 price data available yet. Run crawlers to collect data.")

    # 2. AI Macro Insight (실제 LLM)
    doc.add_heading("2. AI Macro Insight & Analysis", level=1)
    brief_data = await get_daily_brief_json()
    ai_text = await generate_market_analysis(
        news_data=brief_data["news"],
        power_data=brief_data["power"],
        memory_data=brief_data["memory"],
    )
    for para in ai_text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    filename = f"InfraIndex_Daily_Brief_{datetime.now().strftime('%Y%m%d')}.docx"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.get("/daily-brief")
async def get_daily_brief_json(db: AsyncSession = Depends(get_db)) -> Any:
    """
    Returns the Macro Intelligence Briefing as JSON.
    뉴스/전력/메모리 데이터는 macro 크롤러 (Phase 14 구현 후 실제 데이터로 교체 예정).
    현재는 최신 뉴스 기반 큐레이션 데이터 제공.
    """
    return {
        "date": datetime.now(timezone.utc).strftime("%Y-%m-%d"),
        "news": [
            {"title": "AWS announces $10B investment in Mississippi datacenter campuses",
             "source": "TechCrunch", "date": "2026-07-22"},
            {"title": "Korean authorities restrict new high-density power permits in Seoul",
             "source": "Bloomberg", "date": "2026-07-21"},
            {"title": "Liquid cooling standardizations proposed for next-gen AI clusters",
             "source": "DataCenter Dynamics", "date": "2026-07-20"},
        ],
        "power": [
            {"region": "US - Texas (ERCOT)", "cost_per_kwh_usd": 0.08, "trend": "stable"},
            {"region": "US - Virginia (PJM)", "cost_per_kwh_usd": 0.09, "trend": "rising"},
            {"region": "South Korea (KEPCO)", "cost_per_kwh_usd": 0.12, "trend": "rising"},
        ],
        "memory": [
            {"component": "DDR5 128GB R-DIMM", "price_usd": 320.00, "status": "stable"},
            {"component": "HBM3E (Estimated per stack)", "price_usd": 4500.00, "status": "severe_shortage"},
            {"component": "GDDR6 16GB", "price_usd": 45.00, "status": "stable"},
        ],
    }

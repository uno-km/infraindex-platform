import io
import os
from typing import Any, List, Dict
from datetime import datetime, timedelta, timezone
from fastapi import APIRouter, Depends, Query
from fastapi.responses import StreamingResponse, FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func, cast, Date
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor

from shared.db.session import get_db
from apps.server.core.ai_service import generate_market_analysis
from apps.batch.services.gpu.models_history import GpuPriceHistory

router = APIRouter()


async def _fetch_price_data(
    db: AsyncSession, gpu_model_id: str, days: int = 30
) -> List[Dict]:
    """price_history 테이블에서 일별 집계 데이터 조회."""
    since = datetime.now(timezone.utc) - timedelta(days=days)
    day_col = cast(GpuPriceHistory.ts, Date).label("day")

    result = await db.execute(
        select(
            day_col,
            GpuPriceHistory.prv_id,
            GpuPriceHistory.gpu_mdl,
            func.min(GpuPriceHistory.prc_ph).label("min_price"),
            func.avg(GpuPriceHistory.prc_ph).label("avg_price"),
            func.max(GpuPriceHistory.prc_ph).label("max_price"),
            func.count(GpuPriceHistory.id).label("cnt"),
        )
        .where(GpuPriceHistory.gpu_mdl.ilike(f"%{gpu_model_id}%"))
        .where(GpuPriceHistory.ts >= since)
        .group_by(day_col, GpuPriceHistory.prv_id, GpuPriceHistory.gpu_mdl)
        .order_by(day_col.asc(), GpuPriceHistory.prv_id)
    )
    rows = result.all()
    return [
        {
            "date": str(r.day),
            "provider": r.prv_id,
            "gpu_model": r.gpu_mdl,
            "min_price": round(float(r.min_price), 4),
            "avg_price": round(float(r.avg_price), 4),
            "max_price": round(float(r.max_price), 4),
            "data_points": r.cnt,
        }
        for r in rows
    ]


@router.get("/excel")
async def export_excel(
    gpu_model_id: str = Query("H100", description="GPU 모델명"),
    days: int = Query(30, ge=1, le=365, description="조회 기간(일)"),
    db: AsyncSession = Depends(get_db),
) -> Any:
    """
    Exports time-series pricing data as a styled Excel file.
    (수정: Mock 데이터 제거 → 실제 price_history DB 조회)
    """
    # --- File-based Disk Cache (Cost-Effective Caching) ---
    cache_dir = os.path.join(os.getcwd(), "storage", "cache", "excel")
    os.makedirs(cache_dir, exist_ok=True)
    filename = f"{gpu_model_id}_pricing_last{days}days_{datetime.now().strftime('%Y%m%d')}.xlsx"
    cache_file_path = os.path.join(cache_dir, filename)

    if os.path.exists(cache_file_path):
        # Serve from disk cache if generated today
        return FileResponse(
            path=cache_file_path,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    rows = await _fetch_price_data(db, gpu_model_id, days)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"{gpu_model_id} Price Data"

    # ── 헤더 스타일 ──────────────────────────────────────────────
    HEADER_FONT  = Font(bold=True, color="FFFFFF", size=11)
    HEADER_FILL  = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    CENTER       = Alignment(horizontal="center", vertical="center")
    THIN_BORDER  = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    EVEN_FILL    = PatternFill(start_color="EBF3FB", end_color="EBF3FB", fill_type="solid")

    headers = ["Date", "Provider", "GPU Model", "Min $/hr", "Avg $/hr", "Max $/hr", "Data Points"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font      = HEADER_FONT
        cell.fill      = HEADER_FILL
        cell.alignment = CENTER
        cell.border    = THIN_BORDER

    # ── 데이터 삽입 ─────────────────────────────────────────────
    if not rows:
        # 데이터 없음 안내 행
        ws.append(["No data available for the selected period.", "", "", "", "", "", ""])
    else:
        for i, row in enumerate(rows, start=2):
            ws.append([
                row["date"], row["provider"], row["gpu_model"],
                row["min_price"], row["avg_price"], row["max_price"], row["data_points"],
            ])
            if i % 2 == 0:
                for cell in ws[i]:
                    cell.fill   = EVEN_FILL
                    cell.border = THIN_BORDER
            else:
                for cell in ws[i]:
                    cell.border = THIN_BORDER

    # ── 숫자 포맷 ────────────────────────────────────────────────
    for row in ws.iter_rows(min_row=2, min_col=4, max_col=6):
        for cell in row:
            cell.number_format = '"$"#,##0.0000'

    # ── 열 너비 자동 조정 ────────────────────────────────────────
    for col_idx, col in enumerate(ws.columns, start=1):
        max_len = max((len(str(c.value or "")) for c in col), default=10)
        ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 4, 40)

    # ── 메타 시트 ────────────────────────────────────────────────
    ws_meta = wb.create_sheet("Report Info")
    ws_meta.append(["Generated", datetime.now(timezone.utc).isoformat()])
    ws_meta.append(["GPU Model", gpu_model_id])
    ws_meta.append(["Period", f"Last {days} days"])
    ws_meta.append(["Rows", len(rows)])
    ws_meta.append(["Source", "InfraIndex price_history DB"])

    # Save to disk cache
    wb.save(cache_file_path)

    return FileResponse(
        path=cache_file_path,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.get("/word")
async def export_word(db: AsyncSession = Depends(get_db)) -> Any:
    """
    Exports a Daily Briefing as a Microsoft Word Document.
    AI 섹션은 실제 LLM 연결 (ai_service.py) 사용.
    """
    doc = Document()
    doc.add_heading("InfraIndex Daily Macro Briefing", 0)

    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    doc.add_paragraph(f"Report Generated: {date_str}")

    # 1. GPU Price Trend (DB에서 H100 최근 7일)
    doc.add_heading("1. GPU Market Price Trend", level=1)
    h100_rows = await _fetch_price_data(db, "H100", days=7)

    if h100_rows:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Date", "Provider", "Min $/hr", "Avg $/hr", "Max $/hr"]):
            hdr[i].text = h

        for r in h100_rows[:20]:
            cells = table.add_row().cells
            cells[0].text = r["date"]
            cells[1].text = r["provider"]
            cells[2].text = f"${r['min_price']:.4f}"
            cells[3].text = f"${r['avg_price']:.4f}"
            cells[4].text = f"${r['max_price']:.4f}"
    else:
        doc.add_paragraph("No H100 price data available yet. Run crawlers to collect data.")

    # 2. AI Macro Insight (실제 LLM)
    doc.add_heading("2. AI Macro Insight & Analysis", level=1)
    brief_data = await get_daily_brief_json()
    ai_text = await generate_market_analysis(
        news_data=brief_data["news"],
        power_data=brief_data["power"],
        memory_data=brief_data["memory"],
    )
    for para in ai_text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    filename = f"InfraIndex_Daily_Brief_{datetime.now().strftime('%Y%m%d')}.docx"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.get("/daily-brief")
async def get_daily_brief_json(db: AsyncSession = Depends(get_db)) -> Any:
    """
    Returns the Macro Intelligence Briefing as JSON.
    뉴스: NewsArticle DB에서 최신 반도체 관련 기사 조회.
    메모리: AIModelPriceHistory에서 최신 AI 모델 가격 조회.
    전력: 공개 소스 기반 큐레이션 데이터 (KEPCO/EIA 실시간 API가 없어 정적 유지, 갱신 주기 명시).
    """
    from shared.models.news import NewsArticle
    from shared.models.ai_pricing import AIModelMaster, AIModelPriceHistory
    from sqlalchemy import desc

    today_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    # 1. 뉴스: DB에서 최신 반도체 관련 기사 최대 5건
    news_data = []
    try:
        stmt = (
            select(NewsArticle)
            .where(NewsArticle.is_semiconductor_related == True)  # noqa: E712
            .order_by(desc(NewsArticle.published_at))
            .limit(5)
        )
        result = await db.execute(stmt)
        articles = result.scalars().all()
        for a in articles:
            news_data.append({
                "title": a.title,
                "source": a.source_name or "Unknown",
                "date": a.published_at.strftime("%Y-%m-%d") if a.published_at else today_str,
                "url": a.url,
                "category": a.category,
            })
    except Exception:
        pass

    # 뉴스 DB가 비어있으면 빈 배열 (하드코딩 없음)
    if not news_data:
        news_data = []

    # 2. 전력 가격: KEPCO/ERCOT 등 공개 실시간 API 미제공 → 큐레이션 유지 (갱신일 명시)
    # TODO: EIA API (https://api.eia.gov/) 연동 시 실데이터로 교체
    power_data = [
        {"region": "US - Texas (ERCOT)", "cost_per_kwh_usd": 0.078, "trend": "stable",
         "note": "EIA 공시 기준 (2026-07 기준 고정값, 향후 EIA API 연동 예정)"},
        {"region": "US - Virginia (PJM)", "cost_per_kwh_usd": 0.091, "trend": "rising",
         "note": "EIA 공시 기준 (2026-07 기준 고정값)"},
        {"region": "South Korea (KEPCO)", "cost_per_kwh_usd": 0.119, "trend": "rising",
         "note": "KEPCO 산업용 전력단가 기준 (2026-07 기준 고정값)"},
    ]

    # 3. AI 모델 가격 (메모리/AI 시장 지표 대용): OpenRouter 크롤러 수집 데이터
    memory_data = []
    try:
        # 티어별 최신 가격 하나씩
        tier_samples = ["Tier 1", "Tier 2", "Tier 3"]
        for tier in tier_samples:
            stmt = (
                select(AIModelMaster, AIModelPriceHistory)
                .join(AIModelPriceHistory, AIModelMaster.id == AIModelPriceHistory.model_id)
                .where(AIModelMaster.tier == tier)
                .order_by(desc(AIModelPriceHistory.collected_date))
                .limit(1)
            )
            result = await db.execute(stmt)
            row = result.first()
            if row:
                master, history = row
                memory_data.append({
                    "component": f"{master.name} ({master.provider})",
                    "tier": tier,
                    "input_price_1m_usd": history.input_price_1m,
                    "output_price_1m_usd": history.output_price_1m,
                    "collected_date": history.collected_date.isoformat(),
                    "status": "live_data",
                })
    except Exception:
        pass

    return {
        "date": today_str,
        "news": news_data,
        "power": power_data,
        "ai_pricing_snapshot": memory_data,
        "data_sources": {
            "news": "InfraIndex NewsArticle DB (Tier1 RSS 크롤러)",
            "power": "EIA/KEPCO 큐레이션 (고정값 - EIA API 연동 예정)",
            "ai_pricing": "OpenRouter API (실시간 크롤링)",
        }
    }



from shared.models.reporter import DailyReport
from apps.batch.services.reporter.pdf_generator import PDFReporter
from datetime import date
from fastapi import HTTPException

@router.get("/pdf")
async def list_pdf_reports(
    db: AsyncSession = Depends(get_db),
    skip: int = 0,
    limit: int = 100
) -> Any:
    """
    Retrieve list of generated daily PDF reports.
    """
    stmt = select(DailyReport).order_by(DailyReport.report_date.desc(), DailyReport.generated_at.desc()).offset(skip).limit(limit)
    result = await db.execute(stmt)
    reports = result.scalars().all()
    
    return [
        {
            "id": str(r.id),
            "report_date": r.report_date.strftime("%Y-%m-%d"),
            "report_type": r.report_type,
            "file_path": r.file_path,
            "file_size_bytes": r.file_size_bytes,
            "generated_at": r.generated_at.isoformat() if r.generated_at else None
        } for r in reports
    ]

@router.post("/pdf/generate")
async def generate_pdf_report_manually(
    report_type: str = Query("morning", description="morning or evening"),
    target_date: date = Query(None, description="Target date (default today)")
) -> Any:
    """
    Manually trigger PDF report generation.
    """
    if not target_date:
        target_date = date.today()
        
    if report_type not in ["morning", "evening"]:
        raise HTTPException(status_code=400, detail="Invalid report type. Must be 'morning' or 'evening'.")
        
    reporter = PDFReporter()
    try:
        import asyncio
        # Run in a separate thread to avoid event loop conflicts with Playwright
        file_path = await asyncio.to_thread(
            asyncio.run,
            reporter.generate_report(target_date, report_type)
        )
        return {"status": "success", "file_path": file_path, "message": f"Report generated at {file_path}"}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate report: {repr(e)}")

